from __future__ import annotations

import csv
from pathlib import Path

from sirius.config import Settings

_PREVIEW_ROWS = 12
_PREVIEW_CHARS = 2500


class FileService:
    """Per-user file workspace (data/files/<user_id>/) for uploaded datasets.

    CSV/Excel files land here so the Python tool can analyze them; previews
    let the model inspect structure without loading everything into context.
    """

    def __init__(self, settings: Settings) -> None:
        self._root = settings.data_dir / "files"

    def workdir(self, user_id: str) -> Path:
        path = self._root / _safe(user_id)
        path.mkdir(parents=True, exist_ok=True)
        return path

    def save(self, user_id: str, filename: str, content: bytes) -> Path:
        path = self.workdir(user_id) / _safe(filename)
        path.write_bytes(content)
        return path

    def list_files(self, user_id: str) -> list[dict]:
        return [
            {"name": p.name, "size_bytes": p.stat().st_size}
            for p in sorted(self.workdir(user_id).iterdir())
            if p.is_file()
        ]

    def preview(self, user_id: str, filename: str) -> str:
        path = self.workdir(user_id) / _safe(filename)
        if not path.is_file():
            raise FileNotFoundError(f"File not found: {filename}")
        suffix = path.suffix.lower()
        if suffix == ".csv":
            return _preview_csv(path)
        if suffix in (".xlsx", ".xlsm", ".xls"):
            return _preview_excel(path)
        return path.read_text(encoding="utf-8", errors="replace")[:_PREVIEW_CHARS]


def extract_docx_text(path: Path) -> str:
    """Plain text from a Word document (paragraphs + tables)."""
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _preview_csv(path: Path) -> str:
    lines = []
    with path.open(newline="", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for i, row in enumerate(reader):
            if i >= _PREVIEW_ROWS:
                lines.append("[... more rows]")
                break
            lines.append(" | ".join(row))
    return "\n".join(lines)[:_PREVIEW_CHARS]


def _preview_excel(path: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    lines = [f"Sheets: {', '.join(wb.sheetnames)}", f"— first sheet: {wb.sheetnames[0]} —"]
    sheet = wb[wb.sheetnames[0]]
    for i, row in enumerate(sheet.iter_rows(values_only=True)):
        if i >= _PREVIEW_ROWS:
            lines.append("[... more rows]")
            break
        lines.append(" | ".join("" if v is None else str(v) for v in row))
    wb.close()
    return "\n".join(lines)[:_PREVIEW_CHARS]


def _safe(name: str) -> str:
    return Path(name).name.replace("..", "_")
